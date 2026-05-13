"""DOCX 内容提取器"""

from pathlib import Path


class DocxExtractor:
    """Word 文档内容提取"""

    SUPPORTED_EXTENSIONS = {".docx"}

    def extract_text(self, file_path: str | Path) -> str:
        """提取 DOCX 中的文本内容"""
        try:
            from docx import Document
        except ImportError:
            return ""

        try:
            doc = Document(str(file_path))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception:
            return ""

    def extract_metadata(self, file_path: str | Path) -> dict:
        """提取 DOCX 元数据"""
        try:
            from docx import Document
        except ImportError:
            return {}

        try:
            doc = Document(str(file_path))
            props = doc.core_properties
            return {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }
        except Exception:
            return {}
